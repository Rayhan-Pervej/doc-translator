"""
doc-translator: Translate a Japanese folder (names + contents) to English.

Output is always created next to the input folder (same parent), named <folder>_english.
The source folder name can be in Japanese characters.

    Windows:  python translate.py "D:/Projects/日本語文書"
              → output: D:/Projects/日本語文書_english

    macOS:    python translate.py "/Users/rayhan/Desktop/日本語文書"
              → output: /Users/rayhan/Desktop/日本語文書_english

Optional flags:
    --pages N   Only translate first N pages of each PDF (useful for testing)

Setup:
    Copy .env.example to .env and fill in your AWS credentials before running.

Note: On Windows, run with: python -X utf8 translate.py ...
  or set PYTHONUTF8=1 in your environment if paths contain Japanese characters.
"""

import sys
import os
import shutil
import csv
from pathlib import Path
from charset_normalizer import from_path

# Ensure stdout/stderr handle Unicode on all platforms (especially Windows)
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# On Windows, force UTF-8 mode so Japanese characters in argv are preserved
if sys.platform == "win32":
    import ctypes
    ctypes.windll.kernel32.SetConsoleCP(65001)
    ctypes.windll.kernel32.SetConsoleOutputCP(65001)

from translator import translate, translate_name, get_actual_usage


# ── File handlers ──────────────────────────────────────────────────────────────

def translate_txt(src: Path, dst: Path, context: str):
    detection = from_path(src).best()
    encoding = detection.encoding if detection else "utf-8"
    text = src.read_text(encoding=encoding, errors="replace")
    translated = translate(text, context=context)
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(translated, encoding="utf-8")


def translate_csv(src: Path, dst: Path, context: str):
    detection = from_path(src).best()
    encoding = detection.encoding if detection else "utf-8"

    with open(src, encoding=encoding, errors="replace", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)

    translated_rows = []
    for row in rows:
        translated_row = []
        for cell in row:
            if cell.strip():
                translated_row.append(translate(cell, context=context))
            else:
                translated_row.append(cell)
        translated_rows.append(translated_row)

    dst.parent.mkdir(parents=True, exist_ok=True)
    with open(dst, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(translated_rows)


def translate_docx(src: Path, dst: Path, context: str):
    from docx2python import docx2python
    from docx import Document

    # Extract full text structure for context
    extracted = docx2python(str(src))

    # Open original to preserve formatting
    doc = Document(str(src))

    # Translate paragraphs at run level to preserve bold/italic/fonts
    for para in doc.paragraphs:
        if para.text.strip():
            # Collect all run text, translate as one block, redistribute
            full_text = para.text
            translated = translate(full_text, context=context)
            # Replace text in first run, clear the rest
            if para.runs:
                para.runs[0].text = translated
                for run in para.runs[1:]:
                    run.text = ""

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text = para.text
                        translated = translate(full_text, context=context)
                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def _translate_drawing_xml(xml_bytes: bytes, context: str) -> bytes:
    """Translate text runs and set fonts to Arial in a DrawingML XML file.

    Matches by local name (:t, :latin, :ea, :cs) regardless of namespace prefix,
    so files that use a non-standard prefix (rare but valid XML) still work.
    """
    import re

    text = xml_bytes.decode("utf-8")

    # Match any namespace-prefixed <*:t> text node, e.g. <a:t>, <p:t>, <wps:t>
    t_pat = re.compile(r'(<\w+:t(?:\s[^>]*)?>)([^<]*)(</\w+:t>)')

    # Collect unique Japanese text values
    unique = list(dict.fromkeys(
        _xml_unescape(m[1]) for m in t_pat.findall(text)
        if m[1].strip() and _has_japanese(_xml_unescape(m[1]))
    ))
    if unique:
        tmap = _batch_translate_xml_texts(unique, context)
        def replace_t(m):
            plain = _xml_unescape(m.group(2))
            xlated = tmap.get(plain, plain)
            return m.group(1) + _xml_escape(xlated) + m.group(3)
        text = t_pat.sub(replace_t, text)

    # Set fonts to Arial — match any prefix: <a:latin>, <wps:latin>, etc.
    text = re.sub(r'(<\w+:latin\b[^>]*\btypeface=")[^"]*(")', r'\1Arial\2', text)
    text = re.sub(r'(<\w+:ea\b[^>]*\btypeface=")[^"]*(")', r'\1Arial\2', text)
    text = re.sub(r'(<\w+:cs\b[^>]*\btypeface=")[^"]*(")', r'\1Arial\2', text)

    # Replace <*:noAutofit/> with <*:normAutofit/> so translated text (which is
    # longer than Japanese) auto-shrinks to fit the existing shape without moving
    # it — connections, icons, and layout all stay exactly in place.
    text = re.sub(r'<(\w+:)noAutofit/>', r'<\1normAutofit/>', text)

    return text.encode("utf-8")


def _has_japanese(text: str) -> bool:
    return any(0x3000 <= ord(c) <= 0x9FFF or 0xF900 <= ord(c) <= 0xFAFF for c in text)


_BATCH_CHUNK_SIZE = 40  # strings per API call — keeps output well under 8096 tokens


def _translate_chunk(chunk: list[str], context: str) -> dict[str, str]:
    """Translate a single chunk of strings (≤ _BATCH_CHUNK_SIZE). Returns {original: translated}."""
    numbered = "\n".join(f"[{i+1}] {t}" for i, t in enumerate(chunk))
    batch_prompt = (
        f"Translate the Japanese portions of each numbered text segment to English.\n"
        f"Return ONLY the translations, one per line, keeping the same [N] numbering.\n"
        f"Rules:\n"
        f"- Translate ONLY Japanese text. Leave English, numbers, symbols, and formulas unchanged.\n"
        f"- Do NOT add notes or extra text.\n\n"
        f"Context: {context}\n\n"
        f"{numbered}"
    )
    raw = translate(batch_prompt, context="")
    result: dict[str, str] = {}
    for ln in raw.splitlines():
        ln = ln.strip()
        if ln.startswith("[") and "]" in ln:
            try:
                idx = int(ln[1:ln.index("]")])
                val = ln[ln.index("]") + 1:].strip()
                result[chunk[idx - 1]] = val
            except (ValueError, IndexError):
                pass
    return result


def _batch_translate_xml_texts(texts: list[str], context: str) -> dict[str, str]:
    """Batch translate a list of unique strings, return {original: translated}.

    Automatically splits into chunks of _BATCH_CHUNK_SIZE so the API response
    never exceeds the 8096 output-token limit.
    """
    unique = list(dict.fromkeys(t for t in texts if t.strip() and _has_japanese(t)))
    if not unique:
        return {}

    result: dict[str, str] = {}
    for i in range(0, len(unique), _BATCH_CHUNK_SIZE):
        chunk = unique[i:i + _BATCH_CHUNK_SIZE]
        chunk_result = _translate_chunk(chunk, context)
        result.update(chunk_result)
        # Warn if any strings in the chunk came back untranslated
        missing = [t for t in chunk if t not in chunk_result]
        if missing:
            chunk_num = i // _BATCH_CHUNK_SIZE + 1
            print(f"    [warn] {len(missing)} string(s) not returned by API in chunk {chunk_num}, keeping originals")
    return result


def _xml_escape(s: str) -> str:
    """Escape special XML characters in text content."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _xml_unescape(s: str) -> str:
    """Unescape XML entities back to plain text for translation."""
    return s.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"').replace("&apos;", "'")


def _translate_shared_strings(xml_bytes: bytes, context: str) -> bytes:
    """
    Translate sharedStrings.xml using pure regex — no ElementTree serialization,
    so the original XML structure, namespaces, and line endings are preserved exactly.

    Plain <si><t>TEXT</t></si>: replace TEXT.
    Rich <si>...<r><t>A</t></r><r><t>B</t></r>...</si>: join all <t> texts within
      each <si>, translate as one string, put result in first <t>, blank the rest.
    """
    import re

    text = xml_bytes.decode("utf-8")

    # ── 1. Collect all <si>…</si> blocks with their positions ──────────────────
    si_pattern = re.compile(r'<si>(.*?)</si>', re.DOTALL)
    matches = list(si_pattern.finditer(text))

    if not matches:
        return xml_bytes

    # ── 2. For each <si>, determine if plain or rich; collect text to translate ─
    # plain: single <t>…</t> (may have xml:space attr)
    # rich:  one or more <r>…<t>…</t>…</r> runs
    plain_t = re.compile(r'<t(?:\s[^>]*)?>([^<]*)</t>')
    run_pattern = re.compile(r'<r\b[^>]*>(.*?)</r>', re.DOTALL)

    entries = []  # (match, is_rich, raw_xml_text, plain_text)
    for m in matches:
        inner = m.group(1)
        runs = run_pattern.findall(inner)
        if runs:
            # Rich text: join all <t> contents across all runs
            all_t = plain_t.findall(inner)
            raw = "".join(all_t)
            plain = _xml_unescape(raw)
            entries.append((m, True, raw, plain))
        else:
            t_match = plain_t.search(inner)
            if t_match:
                raw = t_match.group(1)
                plain = _xml_unescape(raw)
                entries.append((m, False, raw, plain))

    # ── 3. Batch translate unique Japanese strings (using unescaped plain text) ─
    unique = list(dict.fromkeys(
        e[3] for e in entries if e[3].strip() and _has_japanese(e[3])
    ))
    if not unique:
        return xml_bytes

    tmap = _batch_translate_xml_texts(unique, context)

    # ── 4. Rebuild the XML by replacing <si> blocks in reverse order ───────────
    # Reverse order so positions don't shift as we replace
    result = text
    for m, is_rich, raw_xml, plain in reversed(entries):
        translated_plain = tmap.get(plain, plain)
        if translated_plain == plain:
            continue  # nothing to change
        # XML-escape the translation before inserting into XML
        translated = _xml_escape(translated_plain)

        inner = m.group(1)
        start, end = m.start(), m.end()

        if is_rich:
            # Replace first <t>…</t> content with full translation, blank the rest
            t_positions = list(plain_t.finditer(inner))
            if not t_positions:
                continue
            new_inner = inner
            # Process in reverse so offsets stay valid
            for i, tp in enumerate(reversed(t_positions)):
                new_text = translated if i == len(t_positions) - 1 else ""
                new_inner = (
                    new_inner[:tp.start(1)]
                    + new_text
                    + new_inner[tp.end(1):]
                )
            # Strip <rPh>…</rPh> and <phoneticPr …/> — their sb/eb offsets
            # reference character positions in the original Japanese text; keeping
            # them after translation causes Excel's "repair" warning because the
            # offsets are now out of range for the translated English string.
            new_inner = re.sub(r'<rPh\b[^>]*>.*?</rPh>', '', new_inner, flags=re.DOTALL)
            new_inner = re.sub(r'<phoneticPr\b[^>]*/>', '', new_inner)
            result = result[:start] + f"<si>{new_inner}</si>" + result[end:]
        else:
            # Plain: replace the single <t> content
            tp = plain_t.search(inner)
            if tp:
                new_inner = inner[:tp.start(1)] + translated + inner[tp.end(1):]
                # Strip phonetic runs — invalid after translation (offsets point to Japanese chars)
                new_inner = re.sub(r'<rPh\b[^>]*>.*?</rPh>', '', new_inner, flags=re.DOTALL)
                new_inner = re.sub(r'<phoneticPr\b[^>]*/>', '', new_inner)
                result = result[:start] + f"<si>{new_inner}</si>" + result[end:]

    return result.encode("utf-8")


def _translate_sheet_xml(xml_bytes: bytes, context: str) -> bytes:
    """Translate inline strings and cached formula string values in a worksheet XML."""
    import re
    text = xml_bytes.decode("utf-8")

    # 1. Inline strings: <is><t>...</t></is>
    inline = re.findall(r'(<is><t>)(.*?)(</t></is>)', text, re.DOTALL)
    inline_unique = [t[1] for t in inline if t[1].strip()]

    # 2. Cached formula string results: <c ... t="str"><f ...>...</f><v>CACHED</v></c>
    cached = re.findall(r'(<c [^>]*t="str"[^>]*>(?:<f[^>]*>[^<]*</f>)?<v>)([^<]+)(</v>)', text)
    cached_unique = [t[1] for t in cached if t[1].strip()]

    all_unique = list(dict.fromkeys(inline_unique + cached_unique))
    if not all_unique:
        return xml_bytes

    tmap = _batch_translate_xml_texts(all_unique, context)

    def replace_inline(m):
        orig = m.group(2)
        xlated = tmap.get(_xml_unescape(orig), _xml_unescape(orig))
        return m.group(1) + _xml_escape(xlated) + m.group(3)

    def replace_cached(m):
        orig = m.group(2)
        xlated = tmap.get(_xml_unescape(orig), _xml_unescape(orig))
        return m.group(1) + _xml_escape(xlated) + m.group(3)

    text = re.sub(r'(<is><t>)(.*?)(</t></is>)', replace_inline, text, flags=re.DOTALL)
    text = re.sub(r'(<c [^>]*t="str"[^>]*>(?:<f[^>]*>[^<]*</f>)?<v>)([^<]+)(</v>)', replace_cached, text)
    return text.encode("utf-8")


def _translate_workbook_xml(xml_bytes: bytes, context: str) -> tuple[bytes, dict[str, str]]:
    """Translate <sheet> tab names using regex to avoid ET rewriting namespaces."""
    import re

    # Extract all sheet name= attributes — only from <sheet ...> tags
    names = re.findall(r'(<sheet\b[^>]*\bname=")([^"]+)(")', xml_bytes.decode("utf-8"))
    jp_names = [n[1] for n in names if _has_japanese(n[1])]
    if not jp_names:
        return xml_bytes, {}

    tmap = _batch_translate_xml_texts(jp_names, context)

    # Sanitize and deduplicate translated names in sheet order
    # (process in original sheet order so dedup suffix is deterministic)
    all_names = [n[1] for n in names]  # all sheet names in order (jp + non-jp)
    seen: set[str] = set()
    # Pre-populate seen with non-Japanese names that won't be translated
    for n in all_names:
        if not _has_japanese(n):
            seen.add(n.lower())

    sanitized: dict[str, str] = {}
    for n in all_names:
        if n not in tmap:
            continue
        en = tmap[n]
        # Excel sheet names cannot contain: \ / ? * [ ] :  and max 31 chars
        for ch in r'\/?*[]':
            en = en.replace(ch, "-")
        en = en.replace(":", "-").strip()
        en = en[:31].rstrip()
        # Deduplicate: if name already used, append _2, _3, etc.
        base = en
        suffix = 2
        while en.lower() in seen:
            tag = f"_{suffix}"
            en = base[:31 - len(tag)] + tag
            suffix += 1
        seen.add(en.lower())
        sanitized[n] = en
        tmap[n] = en  # keep tmap in sync for formula patching

    text = xml_bytes.decode("utf-8")
    for jp, en in sanitized.items():
        en_attr = en.replace("&", "&amp;").replace('"', "&quot;")
        text = text.replace(f'name="{jp}"', f'name="{en_attr}"')

    return text.encode("utf-8"), tmap


def _patch_formula_sheet_refs(xml_bytes: bytes, sheet_name_map: dict[str, str]) -> bytes:
    """Replace Japanese sheet name references in formulas, e.g. '変更履歴'!A1 → 'Change History'!A1"""
    if not sheet_name_map:
        return xml_bytes
    text = xml_bytes.decode("utf-8")
    for jp, en in sheet_name_map.items():
        # quoted: 'シート名'! → 'Sheet Name'!
        text = text.replace(f"'{jp}'!", f"'{en}'!")
        # unquoted: シート名! → 'Sheet Name'! (add quotes since English name may have spaces)
        text = text.replace(f"{jp}!", f"'{en}'!")
    return text.encode("utf-8")


def _set_fonts_arial_styles(xml_bytes: bytes) -> bytes:
    """Set all font names in styles.xml to Arial using regex."""
    import re
    text = xml_bytes.decode("utf-8")
    # <name val="MS PGothic"/> → <name val="Arial"/>
    text = re.sub(r'(<name\b[^>]*\bval=")[^"]*(")', r'\1Arial\2', text)
    # <scheme val="minor"/> → <scheme val="none"/>
    text = re.sub(r'(<scheme\b[^>]*\bval=")[^"]*(")', r'\1none\2', text)
    return text.encode("utf-8")


def _set_fonts_arial_shared_strings(xml_bytes: bytes) -> bytes:
    """Set all rFont values in sharedStrings.xml to Arial using regex."""
    import re
    text = xml_bytes.decode("utf-8")
    text = re.sub(r'(<rFont\b[^>]*\bval=")[^"]*(")', r'\1Arial\2', text)
    return text.encode("utf-8")


def _translate_docprops_app(xml_bytes: bytes, sheet_name_map: dict[str, str]) -> bytes:
    """Patch sheet names in docProps/app.xml using the already-translated sheet name map."""
    import re
    text = xml_bytes.decode("utf-8")
    for jp, en in sheet_name_map.items():
        text = text.replace(f">{jp}<", f">{en}<")
    return text.encode("utf-8")


def translate_xlsx(src: Path, dst: Path, context: str):
    import zipfile

    dst.parent.mkdir(parents=True, exist_ok=True)
    tmp = dst.with_suffix(".tmp.xlsx")

    # Write to a temp file first — rename to dst only on full success, so a crash never corrupts output
    try:
        # First pass: read all data and build sheet name map from workbook.xml
        sheet_name_map: dict[str, str] = {}
        all_items: list[tuple] = []
        with zipfile.ZipFile(src, "r") as zin:
            for item in zin.infolist():
                data = zin.read(item.filename)
                all_items.append((item, data))

        # Translate workbook.xml first to get the sheet name map
        for item, data in all_items:
            if item.filename == "xl/workbook.xml":
                print(f"    Translating sheet tab names...")
                translated, sheet_name_map = _translate_workbook_xml(data, context)
                break

        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item, orig_data in all_items:
                fname = item.filename
                data = orig_data

                if fname == "xl/sharedStrings.xml":
                    print(f"    Translating shared strings...")
                    data = _translate_shared_strings(data, context)
                    data = _set_fonts_arial_shared_strings(data)

                elif fname == "xl/workbook.xml":
                    data = translated  # already done above

                elif fname == "xl/styles.xml":
                    print(f"    Setting fonts to Arial in styles...")
                    data = _set_fonts_arial_styles(data)

                elif fname.startswith("xl/worksheets/") and fname.endswith(".xml"):
                    print(f"    Translating worksheet: {fname}")
                    data = _translate_sheet_xml(data, context)
                    data = _patch_formula_sheet_refs(data, sheet_name_map)

                elif (
                    fname.startswith("xl/drawings/")
                    and fname.endswith(".xml")
                    and not fname.endswith(".rels")
                ):
                    print(f"    Translating drawing: {fname}")
                    data = _translate_drawing_xml(data, context)

                elif fname == "docProps/app.xml":
                    data = _translate_docprops_app(data, sheet_name_map)

                if data is orig_data:
                    # Unchanged — write with original ZipInfo (preserves flag_bits, timestamps, etc.)
                    zout.writestr(item, data)
                else:
                    # Changed — use a clean ZipInfo so flag_bits/sizes aren't stale
                    new_info = zipfile.ZipInfo(fname)
                    new_info.compress_type = zipfile.ZIP_DEFLATED
                    zout.writestr(new_info, data)

        # All done — move temp to final destination
        if dst.exists():
            dst.unlink()
        tmp.rename(dst)

    except Exception:
        if tmp.exists():
            tmp.unlink()
        raise


def _sample_bg(pix, bbox, scale):
    """Sample background color from 5 points around the bbox edges."""
    points = [
        (bbox.x0 + 2, bbox.y0 - 4),
        ((bbox.x0 + bbox.x1) / 2, bbox.y0 - 4),
        (bbox.x1 - 2, bbox.y0 - 4),
        (bbox.x0 + 2, bbox.y1 + 4),
        ((bbox.x0 + bbox.x1) / 2, bbox.y1 + 4),
    ]
    samples = []
    for px, py in points:
        sx = min(max(int(px * scale), 0), pix.width - 1)
        sy = min(max(int(py * scale), 0), pix.height - 1)
        p = pix.pixel(sx, sy)
        samples.append((p[0], p[1], p[2]))
    r = sum(s[0] for s in samples) / len(samples)
    g = sum(s[1] for s in samples) / len(samples)
    b = sum(s[2] for s in samples) / len(samples)
    return (r / 255, g / 255, b / 255)


def translate_pdf(src: Path, dst: Path, context: str, max_pages: int = None):
    import fitz  # pymupdf

    # Register pymupdf-fonts so "notos" (Noto Sans) is available for full Unicode coverage
    try:
        import fitz.utils
        fitz.Font("notos")  # will raise if not installed — falls back to helv
        _font = "notos"
    except Exception:
        _font = "helv"

    doc = fitz.open(str(src))
    total_pages = len(doc)
    pages_to_process = min(max_pages, total_pages) if max_pages else total_pages
    print(f"  PDF: {total_pages} pages total, translating {pages_to_process} (font: {_font})")

    for page_num in range(pages_to_process):
        page = doc[page_num]
        print(f"  Page {page_num + 1}/{pages_to_process}...")

        # LINE-level extraction: join all spans in a line → one bbox + one text unit
        blocks = page.get_text("dict")["blocks"]
        spans = []
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_spans = line.get("spans", [])
                if not line_spans:
                    continue
                line_text = "".join(s.get("text", "") for s in line_spans).strip()
                if not line_text:
                    continue
                dominant_size = max(
                    set(s.get("size", 10) for s in line_spans),
                    key=lambda sz: sum(1 for s in line_spans if s.get("size", 10) == sz),
                )
                spans.append({
                    "text": line_text,
                    "bbox": line["bbox"],
                    "size": dominant_size,
                    "color": line_spans[0].get("color", 0),
                })

        if not spans:
            continue

        # One batched API call per page
        numbered = "\n".join(f"[{i+1}] {s['text']}" for i, s in enumerate(spans))
        batch_prompt = (
            f"Translate the Japanese portions of each numbered text segment to English.\n"
            f"Return ONLY the translations, one per line, keeping the same [N] numbering.\n"
            f"Rules:\n"
            f"- Translate ONLY Japanese text. Leave English, numbers, code, and symbols unchanged.\n"
            f"- Preserve all symbols exactly as-is: →, ←, ≦, ≧, ×, ÷, •, ·, °, ±, ©\n"
            f"- Preserve numbers and percentages exactly\n"
            f"- Do NOT add notes, explanations, or parenthetical comments\n"
            f"- Do NOT add any text that wasn't in the original\n"
            f"- Return a clean natural translation only\n\n"
            f"Context: {context} (page {page_num + 1})\n\n"
            f"{numbered}"
        )
        raw = translate(batch_prompt, context="")

        translations = {}
        for ln in raw.splitlines():
            ln = ln.strip()
            if ln.startswith("[") and "]" in ln:
                try:
                    idx = int(ln[1:ln.index("]")])
                    val = ln[ln.index("]") + 1:].strip()
                    translations[idx] = val
                except ValueError:
                    pass

        # Render: redact original text, then insert_htmlbox for full Unicode + auto-scale
        pix = page.get_pixmap(dpi=150)
        scale = pix.width / page.rect.width

        for i, span in enumerate(spans):
            translated = translations.get(i + 1, span["text"])
            bbox = fitz.Rect(span["bbox"])
            font_size = max(span["size"] - 1, 6)

            # Decode original text color from packed int
            color_int = span["color"]
            tr = ((color_int >> 16) & 0xFF) / 255
            tg = ((color_int >> 8) & 0xFF) / 255
            tb = (color_int & 0xFF) / 255
            text_color_hex = "#{:02x}{:02x}{:02x}".format(
                int(tr * 255), int(tg * 255), int(tb * 255)
            )

            bg = _sample_bg(pix, bbox, scale)
            bg_hex = "#{:02x}{:02x}{:02x}".format(
                int(bg[0] * 255), int(bg[1] * 255), int(bg[2] * 255)
            )

            # Redact (erase) original text with background fill
            page.add_redact_annot(bbox, fill=bg)

        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

        # Now insert translated text using insert_htmlbox (auto font + auto scale)
        for i, span in enumerate(spans):
            translated = translations.get(i + 1, span["text"])
            bbox = fitz.Rect(span["bbox"])
            font_size = max(span["size"] - 1, 6)

            color_int = span["color"]
            tr = ((color_int >> 16) & 0xFF) / 255
            tg = ((color_int >> 8) & 0xFF) / 255
            tb = (color_int & 0xFF) / 255
            text_color_hex = "#{:02x}{:02x}{:02x}".format(
                int(tr * 255), int(tg * 255), int(tb * 255)
            )

            html = (
                f'<p style="font-family: sans-serif; font-size: {font_size}pt; '
                f'color: {text_color_hex}; margin: 0; padding: 0;">'
                f'{translated}</p>'
            )
            # scale_low=0.5 allows text to shrink up to 50% before truncating
            # Returns (overflow_lines, scale_used) tuple
            rc = page.insert_htmlbox(bbox, html, scale_low=0.5)
            overflow, scale_used = rc if isinstance(rc, tuple) else (rc, 1.0)
            if overflow > 0 or scale_used < 0.75:
                print(f"    [warn] Page {page_num+1} span {i+1}: overflow={overflow} scale={scale_used:.2f}")

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.subset_fonts()
    doc.ez_save(str(dst))
    doc.close()


# ── Router ─────────────────────────────────────────────────────────────────────

HANDLERS = {
    ".txt": translate_txt,
    ".md":  translate_txt,
    ".csv": translate_csv,
    ".docx": translate_docx,
    ".xlsx": translate_xlsx,
    ".pdf": translate_pdf,
}


def process_file(src: Path, dst: Path, context: str, max_pages: int = None):
    ext = src.suffix.lower()
    handler = HANDLERS.get(ext)
    if handler:
        print(f"  Translating content: {src.name}")
        if ext == ".pdf":
            translate_pdf(src, dst, context, max_pages=max_pages)
        else:
            handler(src, dst, context)
    else:
        # Unsupported format — copy as-is
        print(f"  Copying (unsupported format): {src.name}")
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)


# ── Cost estimation ────────────────────────────────────────────────────────────

# AWS Bedrock Claude Sonnet 4.6 pricing (per million tokens)
_PRICE_INPUT_PER_M  = 3.00
_PRICE_OUTPUT_PER_M = 15.00

def _count_japanese_chars(path: Path) -> int:
    """Count Japanese characters in a file across all supported formats."""
    ext = path.suffix.lower()
    count = 0
    try:
        if ext in (".txt", ".md", ".csv"):
            from charset_normalizer import from_path
            detection = from_path(path).best()
            encoding = detection.encoding if detection else "utf-8"
            text = path.read_text(encoding=encoding, errors="replace")
            count = sum(1 for c in text if 0x3000 <= ord(c) <= 0x9FFF or 0xF900 <= ord(c) <= 0xFAFF)

        elif ext == ".docx":
            import zipfile
            with zipfile.ZipFile(path) as z:
                for name in z.namelist():
                    if name.endswith(".xml"):
                        content = z.read(name).decode("utf-8", errors="replace")
                        count += sum(1 for c in content if 0x3000 <= ord(c) <= 0x9FFF or 0xF900 <= ord(c) <= 0xFAFF)

        elif ext == ".xlsx":
            import zipfile
            with zipfile.ZipFile(path) as z:
                for name in z.namelist():
                    if name.endswith(".xml"):
                        content = z.read(name).decode("utf-8", errors="replace")
                        count += sum(1 for c in content if 0x3000 <= ord(c) <= 0x9FFF or 0xF900 <= ord(c) <= 0xFAFF)

        elif ext == ".pdf":
            import fitz
            doc = fitz.open(str(path))
            for page in doc:
                text = page.get_text()
                count += sum(1 for c in text if 0x3000 <= ord(c) <= 0x9FFF or 0xF900 <= ord(c) <= 0xFAFF)
            doc.close()
    except Exception:
        pass
    return count


def estimate_cost(src: Path) -> None:
    """Scan src (file or folder), estimate token usage and cost, ask user to confirm."""
    print("\n── Estimating cost ──────────────────────────────────────────────")

    files = []
    if src.is_file():
        files = [src]
    else:
        for dirpath, _, filenames in os.walk(src):
            for fname in filenames:
                files.append(Path(dirpath) / fname)

    supported_exts = {".txt", ".md", ".csv", ".docx", ".xlsx", ".pdf"}
    total_jp_chars = 0
    file_breakdown = []

    for f in files:
        ext = f.suffix.lower()
        if ext not in supported_exts:
            continue
        jp = _count_japanese_chars(f)
        if jp > 0:
            file_breakdown.append((f.name, ext, jp))
            total_jp_chars += jp

    if not file_breakdown:
        print("  No Japanese content found — nothing to translate.")
        return

    # Token estimation:
    # Japanese chars → ~1.5 tokens each (kanji/kana compress well)
    # We send each batch with ~200 token prompt overhead per API call
    # Rough API calls: 1 per sheet/page/file
    # Output: English is ~2x longer than Japanese in chars, ~1 token per word ≈ jp_chars * 0.8
    estimated_input_tokens  = int(total_jp_chars * 1.5)
    estimated_output_tokens = int(total_jp_chars * 0.8)

    input_cost  = (estimated_input_tokens  / 1_000_000) * _PRICE_INPUT_PER_M
    output_cost = (estimated_output_tokens / 1_000_000) * _PRICE_OUTPUT_PER_M
    total_cost  = input_cost + output_cost

    # Print breakdown
    print(f"\n  {'File':<50} {'Type':<6} {'JP chars':>9}")
    print(f"  {'-'*50} {'-'*6} {'-'*9}")
    for name, ext, jp in sorted(file_breakdown, key=lambda x: -x[2]):
        print(f"  {name[:50]:<50} {ext:<6} {jp:>9,}")

    print(f"\n  Total Japanese characters : {total_jp_chars:>10,}")
    print(f"  Est. input tokens         : {estimated_input_tokens:>10,}")
    print(f"  Est. output tokens        : {estimated_output_tokens:>10,}")
    print(f"\n  Input cost  (${_PRICE_INPUT_PER_M:.2f}/M tokens) : ${input_cost:.4f}")
    print(f"  Output cost (${_PRICE_OUTPUT_PER_M:.2f}/M tokens) : ${output_cost:.4f}")
    print(f"  ─────────────────────────────────────────")
    print(f"  Estimated total cost      :  ${total_cost:.4f}  (~${total_cost*1.3:.4f} with overhead)")
    print(f"\n  Model: {os.environ.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-6')}")
    print(f"  Note: Estimate is approximate. Actual cost depends on batch sizes and overhead.")
    print()

    answer = input("  Proceed with translation? (y/n): ").strip().lower()
    if answer != "y":
        print("Aborted.")
        sys.exit(0)
    print()


# ── Main pipeline ──────────────────────────────────────────────────────────────

def translate_folder(src_root: Path, dst_root: Path, max_pages: int = None):
    """
    Walk src_root recursively.
    Translate folder names and file names to English.
    Translate file contents for supported formats.
    Write everything to dst_root mirroring the structure.
    """

    # Step 1: Build a translation map for all folder names first
    # This gives us the full English path context before processing files
    print("\n── Phase 1: Translating folder names ──")
    folder_name_map: dict[Path, str] = {}  # original path -> translated name

    for dirpath, dirnames, _ in os.walk(src_root):
        for dirname in dirnames:
            original = Path(dirpath) / dirname
            print(f"  Folder: {dirname}")
            translated = translate_name(dirname, context=f"subfolder of {Path(dirpath).name}")
            folder_name_map[original] = translated
            print(f"    → {translated}")

    # Step 2: Walk and process all files
    print("\n── Phase 2: Translating files ──")

    skipped = []
    errors = []

    for dirpath, dirnames, filenames in os.walk(src_root):
        src_dir = Path(dirpath)

        # Build the translated output path for this directory
        rel_parts = src_dir.relative_to(src_root).parts
        translated_parts = []
        current = src_root
        for part in rel_parts:
            original_sub = current / part
            translated_part = folder_name_map.get(original_sub, part)
            translated_parts.append(translated_part)
            current = original_sub

        dst_dir = dst_root / Path(*translated_parts) if translated_parts else dst_root
        dst_dir.mkdir(parents=True, exist_ok=True)

        # Build context string from the folder path for better translation
        folder_context = " > ".join(translated_parts) if translated_parts else src_root.name

        for filename in filenames:
            src_file = src_dir / filename
            stem = Path(filename).stem
            ext = Path(filename).suffix

            print(f"\nFile: {src_dir.relative_to(src_root) / filename}")

            # Translate file name
            translated_stem = translate_name(stem, context=f"file in folder: {folder_context}")
            translated_filename = translated_stem + ext
            print(f"  Name → {translated_filename}")

            dst_file = dst_dir / translated_filename
            context = f"Document '{translated_filename}' in folder '{folder_context}'"

            try:
                process_file(src_file, dst_file, context, max_pages=max_pages)
            except Exception as e:
                print(f"  ERROR: {e}")
                errors.append((str(src_file), str(e)))
                # Copy original as fallback
                try:
                    shutil.copy2(src_file, dst_dir / filename)
                    print(f"  Fallback: copied original file as {filename}")
                except Exception:
                    pass

    # Summary
    if errors:
        print(f"\n{len(errors)} file(s) had errors:")
        for path, err in errors:
            print(f"  {path}: {err}")
    else:
        print("\nAll files translated successfully.")


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python translate.py <input_folder> [--pages N]")
        print()
        print("Output is always created next to the input folder (same parent directory).")
        print("The output folder is named '<input_folder_name>_english'.")
        print()
        print("Examples (source folder name can be in Japanese):")
        print('  Windows:  python translate.py "D:\\Projects\\日本語文書"')
        print('  macOS:    python translate.py "/Users/rayhan/Desktop/日本語文書"')
        print()
        print("Tip: On Windows use  python -X utf8 translate.py ...  for Japanese folder names.")
        sys.exit(1)

    src = Path(os.fsdecode(sys.argv[1].strip('"').strip("'"))).resolve()

    max_pages = None
    if "--pages" in sys.argv:
        idx = sys.argv.index("--pages")
        max_pages = int(sys.argv[idx + 1])
        print(f"PDF page limit: {max_pages}")

    if not src.exists():
        print(f"Error: input not found: {src}")
        sys.exit(1)

    # ── Single file mode ──────────────────────────────────────────────────────
    if src.is_file():
        print(f"Source:  {src}")
        estimate_cost(src)  # scan + confirm before doing anything
        stem = src.stem
        ext = src.suffix
        translated_stem = translate_name(stem, context=f"file in folder: {src.parent.name}")
        dst_dir = src.parent.parent / (src.parent.name + "_english")
        dst_dir.mkdir(parents=True, exist_ok=True)
        dst_file = dst_dir / (translated_stem + ext)
        print(f"Output:  {dst_file}")
        if dst_file.exists():
            answer = input("Output file already exists. Overwrite? (y/n): ").strip().lower()
            if answer != "y":
                print("Aborted.")
                sys.exit(0)
        context = f"Document '{translated_stem + ext}' in folder '{src.parent.name}'"
        process_file(src, dst_file, context, max_pages=max_pages)
        actual = get_actual_usage()
        print("\n── Done ──")
        print(f"Output: {dst_file}")
        print(f"Actual tokens used  — input: {actual['input_tokens']:,}  output: {actual['output_tokens']:,}")
        print(f"Actual cost         : ${actual['cost_usd']:.4f}")
        sys.exit(0)

    # ── Folder mode ───────────────────────────────────────────────────────────
    dst = src.parent / (src.name + "_english")

    print(f"Source:  {src}")
    print(f"Output:  {dst}")
    estimate_cost(src)  # scan + confirm before doing anything

    if dst.exists():
        answer = input("Output folder already exists. Overwrite? (y/n): ").strip().lower()
        if answer != "y":
            print("Aborted.")
            sys.exit(0)

    dst.mkdir(parents=True, exist_ok=True)
    translate_folder(src, dst, max_pages=max_pages)
    actual = get_actual_usage()
    print("\n── Done ──")
    print(f"Actual tokens used  — input: {actual['input_tokens']:,}  output: {actual['output_tokens']:,}")
    print(f"Actual cost         : ${actual['cost_usd']:.4f}")
