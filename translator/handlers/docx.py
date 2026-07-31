from pathlib import Path

from ..common import has_japanese, batch_translate


def translate_docx(src: Path, dst: Path, context: str):
    from docx import Document

    doc = Document(str(src))

    # collect all unique JP strings from paragraphs and table cells
    texts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip() and has_japanese(para.text):
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() and has_japanese(para.text):
                        texts.append(para.text)

    if not texts:
        dst.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dst))
        return

    tmap = batch_translate(texts, context)

    # apply translations back
    for para in doc.paragraphs:
        if para.text.strip() and para.text in tmap:
            translated = tmap[para.text]
            if para.runs:
                para.runs[0].text = translated
                for run in para.runs[1:]:
                    run.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() and para.text in tmap:
                        translated = tmap[para.text]
                        if para.runs:
                            para.runs[0].text = translated
                            for run in para.runs[1:]:
                                run.text = ""

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
