"""
Cost estimator: scans files and folders to predict token usage before translating.
Never makes any API calls.
"""

import csv
import os
import re
import zipfile
from pathlib import Path

from .common import has_japanese, jp_char_count

_PRICE_INPUT_PER_M  = 3.00
_PRICE_OUTPUT_PER_M = 15.00
SUPPORTED_EXTS      = {".txt", ".md", ".csv", ".docx", ".xlsx", ".pdf"}


def _scan_file(path: Path) -> dict:
    """
    Return how much translatable content is in one file:
      jp_chars: Japanese characters in content
      name_calls: translate_name() calls inside the file (xlsx sheet tabs)
    """
    ext    = path.suffix.lower()
    result = {"jp_chars": 0, "name_calls": 0}
    try:
        if ext in (".txt", ".md"):
            result["jp_chars"] = jp_char_count(path.read_text(encoding="utf-8", errors="replace"))

        elif ext == ".csv":
            from charset_normalizer import from_path
            enc = from_path(path).best()
            enc = enc.encoding if enc else "utf-8"
            with open(path, encoding=enc, errors="replace", newline="") as f:
                rows = list(csv.reader(f))
            result["jp_chars"] = sum(jp_char_count(cell) for row in rows for cell in row)

        elif ext == ".docx":
            from docx import Document
            doc = Document(str(path))
            for para in doc.paragraphs:
                result["jp_chars"] += jp_char_count(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            result["jp_chars"] += jp_char_count(para.text)

        elif ext == ".xlsx":
            with zipfile.ZipFile(path) as z:
                names = z.namelist()

                # sharedStrings: extract only <t> tag text content, not raw XML
                if "xl/sharedStrings.xml" in names:
                    ss = z.read("xl/sharedStrings.xml").decode("utf-8", errors="replace")
                    for t in re.findall(r'<t(?:\s[^>]*)?>([^<]*)</t>', ss):
                        result["jp_chars"] += jp_char_count(t)

                # drawings: discover via .rels files (same logic as handler)
                import posixpath
                drawing_names: set[str] = set()
                drawing_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing"
                for n in names:
                    if n.startswith("xl/worksheets/_rels/") and n.endswith(".rels"):
                        rels = z.read(n).decode("utf-8", errors="replace")
                        for target in re.findall(rf'Type="{re.escape(drawing_type)}"[^>]*Target="([^"]+)"', rels):
                            full = posixpath.normpath(posixpath.join("xl/worksheets/", target)).lstrip("/")
                            drawing_names.add(full)
                if not drawing_names:
                    drawing_names = {n for n in names if n.startswith("xl/drawings/") and n.endswith(".xml") and not n.endswith(".rels")}
                for n in drawing_names:
                    xml = z.read(n).decode("utf-8", errors="replace")
                    for t in re.findall(r'<\w+:t(?:\s[^>]*)?>([^<]*)</\w+:t>', xml):
                        result["jp_chars"] += jp_char_count(t)

                # worksheets: extract inline strings and cached formula string values
                for n in names:
                    if n.startswith("xl/worksheets/") and n.endswith(".xml"):
                        xml = z.read(n).decode("utf-8", errors="replace")
                        for t in re.findall(r'<is><t>([^<]*)</t></is>', xml):
                            result["jp_chars"] += jp_char_count(t)
                        for t in re.findall(r'<c [^>]*t="str"[^>]*>(?:<f[^>]*>[^<]*</f>)?<v>([^<]+)</v>', xml):
                            result["jp_chars"] += jp_char_count(t)

                # sheet tab names
                wb = z.read("xl/workbook.xml").decode("utf-8")
                for tab in re.findall(r'name="([^"]+)"', wb):
                    if has_japanese(tab):
                        result["name_calls"] += 1

        elif ext == ".pdf":
            import fitz
            doc = fitz.open(str(path))
            for page in doc:
                result["jp_chars"] += jp_char_count(page.get_text())
            doc.close()

    except Exception:
        pass  # if a file can't be read, skip it and don't crash the whole estimate
    return result


def estimate_cost(src: Path) -> None:
    """Scan src (file or folder), print cost estimate, ask user to confirm."""
    print("\n── Estimating cost ──────────────────────────────────────────────")

    files: list[Path] = (
        [src] if src.is_file()
        else [Path(dp) / fn for dp, _, fns in os.walk(src) for fn in fns]
    )

    # count how many translate_name() calls will be made (folders + file stems + sheet tabs)
    name_calls = 0
    if src.is_dir():
        seen: set[str] = set()  # deduplicate: same folder name in multiple places costs only one call
        for dirpath, dirnames, filenames in os.walk(src):
            for d in dirnames:
                if has_japanese(d) and d not in seen:
                    seen.add(d)
                    name_calls += 1
            for fn in filenames:
                stem = Path(fn).stem
                if has_japanese(stem) and stem not in seen:
                    seen.add(stem)
                    name_calls += 1

    total_jp_chars = 0
    file_breakdown: list[tuple[str, str, int]] = []

    for f in files:
        if f.suffix.lower() not in SUPPORTED_EXTS:
            continue
        info = _scan_file(f)
        name_calls    += info["name_calls"]
        if info["jp_chars"] > 0:
            file_breakdown.append((f.name, f.suffix.lower(), info["jp_chars"]))
            total_jp_chars += info["jp_chars"]

    if not file_breakdown and name_calls == 0:
        print("  No Japanese content found, nothing to translate.")
        return

    # Multipliers calibrated from real run: 11,414 JP chars produced 21,741 input / 11,197 output tokens.
    # Input: 1.90x JP chars (prompt overhead + numbered list format, no cache discount applied).
    # Output: 1.0x JP chars (index-based output, EN translation only, no JP keys repeated).
    est_input  = int(total_jp_chars * 1.9) + name_calls * 300
    est_output = int(total_jp_chars * 1.0) + name_calls * 50
    in_cost    = est_input  / 1_000_000 * _PRICE_INPUT_PER_M
    out_cost   = est_output / 1_000_000 * _PRICE_OUTPUT_PER_M
    total_cost = in_cost + out_cost

    print(f"\n  {'File':<50} {'Type':<6} {'JP chars':>9}")
    print(f"  {'-'*50} {'-'*6} {'-'*9}")
    for name, ext, jp in sorted(file_breakdown, key=lambda x: -x[2]):
        print(f"  {name[:50]:<50} {ext:<6} {jp:>9,}")

    print(f"\n  Total Japanese characters : {total_jp_chars:>10,}")
    print(f"  Name translation calls    : {name_calls:>10,}  (folders, files, sheet tabs)")
    print(f"  Est. input tokens         : {est_input:>10,}")
    print(f"  Est. output tokens        : {est_output:>10,}")
    print(f"\n  Input cost  (${_PRICE_INPUT_PER_M:.2f}/M tokens) : ${in_cost:.4f}")
    print(f"  Output cost (${_PRICE_OUTPUT_PER_M:.2f}/M tokens) : ${out_cost:.4f}")
    print(f"  ─────────────────────────────────────────")
    print(f"  Estimated total cost      :  ${total_cost:.4f}  (~${total_cost*1.3:.4f} with overhead)")
    print(f"\n  Model: {os.environ.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-6')}")
    print(f"  Note: Multipliers calibrated from real runs. Input: 1.9x JP chars. Output: 1.0x JP chars (index-based, EN only).")
    print()
